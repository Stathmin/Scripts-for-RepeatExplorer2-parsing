import glob as gl
import pandas as pd
import re as re
import itertools as it
import xlsxwriter as xl
import textwrap as txtw
import numpy as np
import os
import plotly.express as px

from collections import defaultdict

import docx
from blast import apply_leven as lev
from blast import blaster

##ПАРСИНГИ

def get_cyphers():
    with open('cypher.csv', 'r') as file:
        #Шифровки для образцов
        cypherdict = defaultdict(None)
        content = {x.split(',')[0]:x.split(',')[1:] for x in file.read().split('\n')}
        cypherdict.update(content)
        return cypherdict


def get_tareans(path, index):
    repeats_paths = gl.glob(f'{path}/*TAREAN*')
    names, seqs = [], []
    for repeat_path in repeats_paths:
        with open(repeat_path, 'r') as file:
            splited_fastas = [*filter(None, file.read().split('>'))]
            for fasta in splited_fastas:
                lines = [*filter(None, fasta.split('\n'))]
                names.append(index+'_'+lines[0])
                seqs.append('\n'.join(txtw.wrap("".join(lines[1:]), width=80)))

    cls = list(map(lambda x: re.findall('CL\d+',x)[0], names))
    cls_zeroes = list(map(lambda x: "CL{:04d}".format(int(re.findall('\d+',x)[0])), cls))
    pics = [gl.glob(f'{path}/seqclust/clustering/clusters/dir_{x}/graph_layout.png')[0] for x in cls_zeroes]
    tareandf = pd.DataFrame({'cluster':names, 'seq':seqs, 'pic_path':pics, 'Cluster':cls})
    return tareandf


def get_copies(path):
    cltab = gl.glob(f'{path}/CLUSTER_TABLE.csv')[0]
    copydf = pd.read_csv(cltab, skiprows=5, delimiter='\t')
    
    with open(cltab, 'r') as file:
        lines=file.readlines()
        total_size=int(lines[4].split('\t')[1].strip())
        
    copydf['Cluster']=copydf['Cluster'].apply(lambda x: 'CL'+str(x))
    copydf['size, %']=copydf['Size'].apply(lambda x: 100*x/total_size)
    return copydf

def get_local_blasts():
    localblastdf = pd.read_csv('TestComponentFeature.csv', delimiter='\t', names=['q','s','prcnt', 'len', 'errs', 'gaps', 'q_st', 'q_end', 's_st', 's_end', 'E', 'score'],  index_col=False)
    IDS = set(localblastdf['q'].to_list())
    localblastdf['BlastType'] = localblastdf['s'].apply(lambda x: 'self' if x in IDS else 'ref')
    return localblastdf

def get_comparative_data(index, path):
    df = pd.read_csv(f'{path}/COMPARATIVE_ANALYSIS_COUNTS.csv', skiprows=2, delimiter='\t', nrows=500)
    local_indexi = df.columns.drop(['cluster','supercluster']).to_list()
    for local_index in local_indexi:
        df[f'{local_index}, %'] = (100*df[local_index]/(df[local_indexi].sum(axis=1))).map('{:,.2f}'.format)
    df['merger'] = df['cluster'].apply(lambda x: index + '_CL' + str(x))
    return df.drop(['cluster', 'supercluster'], axis=1)

def create_doc(index, annotation_db, my_df, ncbi_naming):
    txt = []
    #solos = list(set(annotation_db['кластер'].unique()) - set(my_df['qseqid'].unique()))
    solos=list()
    dopplers = set()

    counts = annotation_db.groupby('TAREAN_annotation').count()['кластер'].rename({'TAREAN_annotation':index}, axis=0)
    doc.add_paragraph(f'{index} {" ".join(get_cyphers()[index][0:2])}\n')
    doc.add_paragraph('Кластерный анализ с использованием программы RepeatExplorer2 позволил выявить:')
    table_writer(doc,pd.DataFrame(counts).reset_index())
    doc.add_paragraph('\nНиже преведены найденные с помощью RepeatExplorer2 повторы, их длины, копийность(отношение числа ридов в кластере к количеству ридов, отобранных для анализа).')
    for type_gr, group in annotation_db.groupby('TAREAN_annotation'):
        doc.add_paragraph('\n'+type_gr)
        group['размер, %'] = group['размер, %'].apply(lambda x:'{0:.3f}'.format(x))
        table_writer(doc, group[['кластер', 'размер, %']])
    doc.add_paragraph('\nПоиск и оценка повторов проводилась с помощью пакета программ BLAST NCBI. Анализ показал, что часть найденных повторов имеет гомологию к ранее известным повторам.')
    
    my_df = my_df.sort_values('num')
    chosen_items = sorted(list(annotation_db['num'].unique()))
    for num in chosen_items:
        local_df = my_df[my_df['num']==num]
        if len(local_df) == 0:
            name = annotation_db[annotation_db['num']==num]['кластер'].values[0]
            solos.append(name)
            doc.add_paragraph('\n'+name)
            doc.add_paragraph('---нет совпадений---')
        else:
            for qseqid, group in local_df.groupby('qseqid'):
                doc.add_paragraph('\n'+qseqid)
                group = group.sort_values(['task', 'fitting', 'sseqid'])[group['sum_coverage']>0.3]
                group['db'] = 'ncbi'
                if len(group) > 0: 
                    outer = lev(group).head()[['qseqid', 'sseqid', 'task', 'sum_coverage']]
                    out = outer.merge(ncbi_naming, left_on='sseqid', right_on='short').drop('sseqid', axis=1)[['full', 'task', 'sum_coverage']]
                    out['sum_coverage'] = out['sum_coverage'].apply(lambda x:'{0:.3f}'.format(x))
                    out = out.rename({'full':'повтор из базы NCBI','sum_coverage':'суммарное покрытие'}, axis=1)
                    table_writer(doc, out)
                    loc_dopplers = out['повтор из базы NCBI'].apply(lambda x: x.split(' ')[1]).unique()
                    dopplers.update(loc_dopplers)
                else:
                    doc.add_paragraph('---нет совпадений---')
                    print(qseqid+'!!!')
                    solos.append(qseqid)
    doc.add_paragraph(f'\nСреди гомологчных повторов встречаются повторы родов {", ".join(sorted([x for x in dopplers if "." not in x]))}.')
    if len(solos) > 0:
        rep_fin_num = [*range(10)]
        rep_incl = ["повторов", "повтор", "повтора", "повтора", "повтора", "повторов", "повторов", "повторов", "повторов", "повторов"]
        rep_incl_dict = dict(zip(rep_fin_num, rep_incl))
        doc.add_paragraph(f'{len(solos)} {rep_incl_dict[len(solos)%10]} не показали гомологий: {", ".join(solos)}')
    doc.save(f'./{index}.docx')
    
def table_writer(doc, df):
    t = doc.add_table(df.shape[0]+1, df.shape[1])

    # add the header rows.
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]

    # add the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])

my_list = ['KA21','KK1','KK8','KP22','KP37','KA22','KK9','KP23','KP38','KA23','KK10','KP1','KP24','KP39','KA1','KA24','KK11','KP10','KP25','KP4','KA10','KA25','KK17','KP11','KP26','KP5','KA11','KA26','KK18','KP12','KP27','KP6','KA12','KA27','KK19','KP13','KP28','KP7','KA13','KA28','KK2','KP14','KP29','KP8','KA14','KA29','KP15','KP3','KP9','KA15','KA3','KK20','KP16','KP30','V1','KA16','KA4','KK21','KP17','KP31','V4','KA17','KA5','KK3','KP18','KP32','V5','KA18','KA6','KP19','KP33','V7','KA19','KA7','KK4','KP2','KP34','KA2','KA8','KP20','KP35','KA20','KA9','KK6','KP21','KP36']

my_list = my_list + ['KA'+str(x+1) for x in range(21)] + ['KA27', 'KA28', 'KA29']


for i in my_list:
    indexi = [i]
    paths = [f'/home/dulyanov/RepexScreens/{x}' for x in indexi]

    my_df = blaster(dbs=['local'], subjects=indexi, tasks=['megablast', 'dc-megablast', 'blastn'], verbose=False)
    my_df = my_df[my_df['db'].apply(lambda x: 'ncbi' in x)]
    my_df['num'] = my_df['qseqid'].apply(lambda x: int(re.findall('CL\d+',x)[0].replace('CL','')))
    

    ind_path_dict = dict(zip(indexi, paths))

    dataframe_list = list()

    with open('./ncbi_naming.csv','r') as file:
        ncbi_naming = pd.read_csv(file).drop('Unnamed: 0', axis=1)

    for index, path in ind_path_dict.items():
        print(index, path)
        df_tar = get_tareans(path, index)
        df_cop = get_copies(path)
        df = pd.merge(df_tar, df_cop, left_on='Cluster', right_on='Cluster').drop(['Cluster', 'Final_annotation'], axis=1)
        dataframe_list.append(df)

    annotation_db = pd.concat(dataframe_list)
    if 'comparative' in annotation_db.columns.values.tolist():
        annotation_db['comparative'] = annotation_db['comparative'].fillna(False)
    annotation_db = annotation_db[annotation_db['TAREAN_annotation'] != 'rDNA']
    annotation_db = annotation_db.rename({'cluster':'кластер', 'size, %':'размер, %'}, axis=1)
    renamer = {'Putative satellites (low confidence)':'Сателлиты (низковероятные)', 'Putative LTR elements':'Вероятные LTR элементы', 'Putative satellites (high confidence)':'Сателлиты (высоковероятные)'}
    annotation_db['TAREAN_annotation'] = annotation_db['TAREAN_annotation'].apply(lambda x: renamer[x])
    annotation_db['num'] = annotation_db['кластер'].apply(lambda x: int(re.findall('CL\d+',x)[0].replace('CL','')))
    
    allowed = set(annotation_db['кластер'].unique())
    my_df = my_df[my_df['qseqid'].apply(lambda x: x in allowed)]

    doc = docx.Document("./blank_word.docx")

    index = index.replace('(old)','')

    create_doc(index, annotation_db, my_df, ncbi_naming)

given_tareans = ['Putative satellites (high confidence)',
 'Putative satellites (low confidence)',
 'Putative LTR elements',
 'rDNA']
our_tareans = ['PShc', 'PSlc', 'pLTR', 'rDNA']

tarean_renamer_dict = dict(zip(given_tareans, our_tareans))
tarean_renamer = lambda x: tarean_renamer_dict[x]
def_type = lambda x: 'solo' if len(re.findall('\d+', x)) == 1 else 'comp'
def_index = lambda x: x.split('_')[0]

total_table = pd.merge(blastdf, annotation_db, left_on='qseqid', right_on='cluster').drop('cluster', axis=1)
total_table['TAREAN_annotation'] = pd.Categorical(total_table['TAREAN_annotation'], given_tareans)
total_table['index'] = pd.Categorical(total_table['qseqid'].apply(def_index), indexi)
total_table['cl_number'] = total_table['qseqid'].apply(lambda x: int(re.findall('CL\d+', x)[0].replace('CL',''))).astype(int)
total_table['db_source'] = total_table['db'].apply(lambda x: x.replace('_x3', '')).replace('comp', 'local')
total_table['db_source'] = pd.Categorical(total_table['db_source'], ['local', 'ref', 'ncbi', 'other'])
total_table.loc[(total_table['sseqid'].apply(def_index).apply(lambda x: x not in indexi)) & (total_table['db_source']=='local'), 'db_source'] = 'other'


total_table = total_table.sort_values(['index', 'TAREAN_annotation', 'cl_number', 'db_source', 'sseqid'])

# indexi = ['KA25', 'KA26', 'V1', 'KA25KA26(1to1)', 'KA25KA26(1to5)']

def set_starting_positions(indexi, blank_cols = 6, blank_rows = 2, entry_length = {'comp':12,'solo':8}):
    cols = dict()
    for num, item in enumerate(indexi):
        result = sum(list(map(lambda x: entry_length[entry_type[x]], indexi[0:num]))) + blank_cols
        cols[item] = result
    cols['ref'] = 0
    cols['ncbi'] = 0
    rows = {key:blank_rows for key in cols.keys()}
    return rows, cols


def clustermerger(worksheet, first_row_of_cluster, col, row, prev_name, prev_size, prev_seq, prev_length, merge_format, merge_format_seq):
    worksheet.merge_range(first_row_of_cluster, col, row-1, col, prev_name, merge_format)
    worksheet.merge_range(first_row_of_cluster, col+1, row-1, col+1, prev_size, merge_format)
    worksheet.merge_range(first_row_of_cluster, col+2, row-1, col+2, prev_seq, merge_format_seq)
    worksheet.merge_range(first_row_of_cluster, col+3, row-1, col+3, prev_length, merge_format)
    worksheet.merge_range(first_row_of_cluster, col+4, row-1, col+4, None, merge_format_seq)



def rewrite_df(row1, row2):
    row1 = pd.DataFrame(row1)
    row2 = pd.DataFrame(row2)
    for col in set(table.columns).intersection(set(annotation_db.columns)):
        #print(col,'\n', row1[col],'\n', row2[col], '\n',)
        row1[col] = row2[col].tolist()[0]
    return row1


def write_to_excel(path=f'./sicktest.xlsx'):
    workbook = xl.Workbook(path)
    worksheet = workbook.add_worksheet('main_sheet')
    
    
    def_type = lambda x: 'solo' if len(re.findall('\d+', x)) == 1 else 'comp'
    entry_type = {index:def_type(index) for index in indexi}
    
    rows, cols = set_starting_positions(indexi)
    
    prev_cls=None
    not_found = []
    
    for num, each in enumerate(head):
        worksheet.write(1,num,each)
    
    return rows, cols

def introduce(worksheet, df, row, col, secondary = False, ignore_picture = False, renaming_db=[]):
    to_drop = ['seq', 'qseqid','sseqid', 'pident', 'length', 'mismatch', 'gapopen', 'qlength', 'fitting', 'sum_coverage', 'sum_lengths', 'source', 'qstart',
       'qend', 'sstart', 'send', 'evalue', 'bitscore', 'task', 'db', 'Supercluster', 'Size', 'Size_adjusted', 'Automatic_annotation',
       'comparative', 'index', 'cl_number', 'db_source'] + indexi
    if secondary:
        to_leave = ['sseqid', 'fitting', 'sum_coverage', 'task', 'db']
    else:
        to_leave = ['qseqid']
    for i in to_leave:
        to_drop.remove(i)
    
    if ignore_picture:
        to_drop.append('pic_path')
    
    needed_df = df.drop(to_drop, axis=1, errors='ignore').drop_duplicates()
    
    needed_df = needed_df.dropna(axis=1, how='all')
    
    order = ['qseqid','sseqid','TAREAN_annotation','size, %','pic_path','KA25, %', 'KA26, %','task','fitting','db','sum_coverage']

    needed_df.reindex(order, axis=1).dropna(axis=1)
    
    fitting_format = {'weak':{'bold': False, 'italic': True},
                       'partial':{'bold': False, 'italic': False},
                       'composite':{'bold': True, 'italic': False},
                       'near-full':{'bold': True, 'italic': False}}
    db_format = lambda x: {'underline': True} if '_x3' in x else {'underline': False}
    task_format = {'blastn':{'font_color':'red'}, 'dc-megablast':{'font_color':'blue'}, 'megablast':{'font_color':'black'}}
    
    if secondary:
        format_params = needed_df[['task','db','fitting']].values.tolist()[0]
        needed_df = needed_df.drop(['task','db','fitting'], axis=1)
        my_format = {**task_format[format_params[0]], **db_format(format_params[1]), **fitting_format[format_params[2]]}
#         my_format = {**fitting_format[format_params[2]]}
        my_format = workbook.add_format(my_format)
    
    for num, column in enumerate(needed_df.columns):
        item = needed_df[column].values[0]
        if len(renaming_db) > 0:
            if (item in renaming_db['short'].values.tolist()):
                item = renaming_db[renaming_db['short']==item]['full'].values[0]
        if column != 'pic_path':
            if column == 'TAREAN_annotation':
                item = tarean_renamer(item)
            if secondary:
                worksheet.write(row, col+num, item, my_format)
            else:
                worksheet.write(row, col+num, item)
        else:
            worksheet.insert_image(row, col+num, item, {'x_scale': 0.05, 'y_scale': 0.05, 'object_position': 1})
        worksheet.write(1,col+num, column)
    return needed_df

def levenshtein(seq1, seq2):
    size_x = len(seq1) + 1
    size_y = len(seq2) + 1
    matrix = np.zeros ((size_x, size_y))
    for x in range(size_x):
        matrix [x, 0] = x
    for y in range(size_y):
        matrix [0, y] = y

    for x in range(1, size_x):
        for y in range(1, size_y):
            if seq1[x-1] == seq2[y-1]:
                matrix [x,y] = min(
                    matrix[x-1, y] + 1,
                    matrix[x-1, y-1],
                    matrix[x, y-1] + 1
                )
            else:
                matrix [x,y] = min(
                    matrix[x-1,y] + 1,
                    matrix[x-1,y-1] + 1,
                    matrix[x,y-1] + 1
                )
                
    return (matrix[size_x - 1, size_y - 1])

def apply_leven(df, renaming_db):
    framelist = []
    print('!!!!!!!!!!!!!')
    print(df)
    for cl, frame in df.groupby('qseqid'):
        
        if len(pd.merge(frame, renaming_db, left_on='sseqid', right_on='short')) > 0:
            frame = pd.merge(frame, renaming_db, left_on='sseqid', right_on='short')
            ncbi_names = sorted(list(set(frame['full'].tolist())))
            list_to_remain=[ncbi_names[0], ]
            for i in ncbi_names[1:]:
                #print(levenshtein(list_to_remain[-1],i))
                if levenshtein(list_to_remain[-1], i) > 0.5*len(min(list_to_remain)):
                    list_to_remain.append(i)
            print('list', list_to_remain)
            frame = frame[frame['full'].isin(list_to_remain)]
            print('final_frame', frame)
            framelist.append(frame)
        else:
            framelist = [frame]
    
    print(len(framelist))
    final_df = pd.concat(framelist)
    return final_df.drop(['short', 'full'], axis=1)

colnames = pd.read_csv('colnames.csv').columns.values.tolist()

with open('./ncbi_naming.csv','r') as file:
    ncbi_naming = pd.read_csv(file).drop('Unnamed: 0', axis=1)

renaming_db = ncbi_naming


workbook = xl.Workbook('worst_job.xlsx')
worksheet = workbook.add_worksheet('main_sheet')

entry_length = {'comp':7,'solo':5}

entry_type = {index:def_type(index) for index in indexi}

rows, cols = set_starting_positions(indexi, entry_length=entry_length)
cols['ref'] = max(cols.values()) + entry_length[entry_type[list(cols.keys())[-3]]]
cols['ncbi'] = cols['ref'] + 1

lefts = [0] + sorted(cols.values())
rights = [coord-1 for coord in lefts[1:]] + [lefts[-1]+1]
merges = list(zip(lefts, rights))

for num, (l,r) in enumerate(merges):
    if l-r !=0:
        worksheet.merge_range(0, l, 0, r, colnames[num])
    else:
        worksheet.write(0, l, colnames[num])
    

prev_cls=None
not_found = []
#TAREAN_annotation 
for qseqid in total_table['qseqid'].drop_duplicates():
    
    row = max(rows.values())
    col = 0
    
    start_row = row
    
    current_df = pd.DataFrame(total_table[total_table['qseqid'] == qseqid])
    if 'rDNA' in current_df['TAREAN_annotation'].values.tolist():
        continue
    needed_df = introduce(worksheet, current_df, row, col)
    db_sources = ['local', 'ref', 'ncbi']
    for db_source in db_sources:
        sub_df = current_df[current_df['db_source']==db_source]
        if (db_source == 'ncbi') & (len(sub_df)>0):
            print('levin')
            sub_df = apply_leven(sub_df, renaming_db)
        if len(sub_df)>0:
            if (('KK9' in sub_df['qseqid'].unique()[0]) | ('KK11' in sub_df['qseqid'].unique()[0])):
                sub_df.to_csv(f'./{qseqid}_{db_source}.chosencsv')
            sub_df['task'] = pd.Categorical(sub_df['task'], ['megablast', 'dc-megablast', 'blastn'])
            sub_df['fitting'] = pd.Categorical(sub_df['fitting'], ['near-full', 'composite', 'partial', 'weak'])
            sub_df = sub_df.sort_values(['task', 'fitting'])
            if db_source == 'local':
                coord_index = lambda x: def_index(x['sseqid'].values[0])
            else:
                coord_index = lambda x: db_source
            #for (_, sseqid), table in sub_df.groupby(['qseqid','sseqid']):
            for _, row in sub_df.iterrows():
                table = row.to_frame().transpose()
                sseqid = table['sseqid'].values[0]
                sseqid_an_df = annotation_db[annotation_db['cluster']==sseqid]
                if len(sseqid_an_df) != 0:
                    table = rewrite_df(table, sseqid_an_df)
                index = coord_index(table)
                col = cols[index]
                row = rows[index]
                if db_source in ['ncbi', 'ref']:
                    table = table[['sseqid','task','db','fitting']]
                needed_df = introduce(worksheet, table, row, col, secondary = True, ignore_picture = db_source in ['ncbi', 'ref'], renaming_db=renaming_db)
                if db_source == 'local':
                    print(index, col)
                rows[index] +=1
                
            
            
    if max(rows.values()) - start_row < 4:
        diff = 4 - max(rows.values()) + start_row
    else:
        diff = 0
    rows = dict(list((key,max(rows.values())+diff) for key, value in rows.items()))

workbook.close()


import glob
import os
import pandas as pd
import re

items = list(map(lambda x: x.split('/')[1], glob.glob('./KA25/index.html')))
misc = ['organelle',
        '|--plastid',
        '|--mitochondria',
        'Unclassified repeat (No evidence)',
        'contamination',]
        
outdfs = []
for path, item in zip([glob.glob(f'./{x}/summarized_annotation.html')[0] for x in items], items):
    df = pd.read_fwf(path, skiprows=13, widths=[49, 16, 17, 12, 30])
    df.columns = ['type', 'percent', 'scl', 'cl', 'reads']
    df = df.drop([74, 75, 79, 80, 82, 83])

    for column in ['percent', 'scl', 'cl', 'reads']:
        df[column] = df[column].apply(lambda x: float(x.replace(' ','').replace('|','').replace('</pre>','')))

    df['type'] = df['type'].apply(lambda x: re.sub(r'[ ]+$', '', ''.join(x[0:-1])).replace("'","|"))
    df = df[['type', 'percent']]
    df.columns = ['type', item]
    df = df.set_index('type')
    not_clusters = df.loc[misc][item].sum()
    factor = (100 - not_clusters)/100
    df[item+' normalized'] = df[item].apply(lambda x: round(x*factor,2))
    df.loc[misc, item+' normalized'] = ''

    
    outdfs.append(df)

df = outdfs[0]
if len(outdfs) > 1:
    for other_df in outdfs[1:]:
        df = pd.merge(left=df, right=other_df, left_on='type', right_on='type')

#df.to_csv('chto.csv')


import glob as gl
import pandas as pd
import re as re
import itertools as it
import xlsxwriter as xl
import textwrap as txtw
import numpy as np
import os
import plotly.express as px

from collections import defaultdict

depth_dict=dict((
("Unclassified_repeat (conflicting evidences)",0),
("|--rDNA",1),
("|   |--45S_rDNA",2),
("|   |   |--18S_rDNA",3),
("|   |   |--25S_rDNA",3),
("|   |   '--5.8S_rDNA",3),
("|   '--5S_rDNA",2),
("|--satellite",1),
("'--mobile_element",1),
("|--Class_I",2),
("|   |--SINE",3),
("|   |--LTR",3),
("|   |   |--Ty1_copia",4),
("|   |   |   |--Ale",5),
("|   |   |   |--Alesia",5),
("|   |   |   |--Angela",5),
("|   |   |   |--Bianca",5),
("|   |   |   |--Bryco",5),
("|   |   |   |--Lyco",5),
("|   |   |   |--Gymco-III",5),
("|   |   |   |--Gymco-I",5),
("|   |   |   |--Gymco-II",5),
("|   |   |   |--Ikeros",5),
("|   |   |   |--Ivana",5),
("|   |   |   |--Gymco-IV",5),
("|   |   |   |--Osser",5),
("|   |   |   |--SIRE",5),
("|   |   |   |--TAR",5),
("|   |   |   |--Tork",5),
("|   |   |   '--Ty1-outgroup",5),
("|   |   '--Ty3_gypsy",4),
("|   |       |--non-chromovirus",5),
("|   |       |   |--non-chromo-outgroup",6),
("|   |       |   |--Phygy",6),
("|   |       |   |--Selgy",6),
("|   |       |   '--OTA",6),
("|   |       |       |--Athila",7),
("|   |       |       '--Tat",7),
("|   |       |           |--TatI",8),
("|   |       |           |--TatII",8),
("|   |       |           |--TatIII",8),
("|   |       |           |--Ogre",8),
("|   |       |           '--Retand",8),
("|   |       '--chromovirus",5),
("|   |           |--Chlamyvir",6),
("|   |           |--Tcn1",6),
("|   |           |--chromo-outgroup",6),
("|   |           |--CRM",6),
("|   |           |--Galadriel",6),
("|   |           |--Tekay",6),
("|   |           |--Reina",6),
("|   |           '--chromo-unclass",6),
("|   |--pararetrovirus",3),
("|   |--DIRS",3),
("|   |--Penelope",3),
("|   '--LINE",3),
("'--Class_II",2),
("|--Subclass_1",3),
("|   '--TIR",4),
("|       |--MITE",5),
("|       |--EnSpm_CACTA",5),
("|       |--hAT",5),
("|       |--Kolobok",5),
("|       |--Merlin",5),
("|       |--MuDR_Mutator",5),
("|       |--Novosib",5),
("|       |--P",5),
("|       |--PIF_Harbinger",5),
("|       |--PiggyBac",5),
("|       |--Sola1",5),
("|       |--Sola2",5),
("|       '--Tc1_Mariner",5),
("'--Subclass_2",3),
("'--Helitron",4),
("organelle",0),
("|--plastid",1),
("'--mitochondria",1),
("Unclassified repeat (No evidence)",0),
("contamination",0),
))


#Для Кати

import glob
import os
import pandas as pd
import re

INDEX='KA26'

items = list(map(lambda x: x.split('/')[1], glob.glob(f'./{INDEX}/index.html')))
misc = ['organelle                                       |',
       ' |--plastid                                     |',
       " '--mitochondria                                |",
       'Unclassified repeat (No evidence)               |',
       'contamination                                   |']
        
outdfs = []
for path, item in zip([glob.glob(f'./{x}/summarized_annotation.html')[0] for x in items], items):
    df = pd.read_fwf(path, skiprows=13, widths=[49, 16, 17, 12, 30], delimiter="\n\t")
    df.columns = ['type', 'percent', 'scl', 'cl', 'reads']
    df = df.drop([74, 75, 79, 80, 82, 83])

    for column in ['percent', 'scl', 'cl', 'reads']:
        df[column] = df[column].apply(lambda x: float(x.replace(' ','').replace('|','').replace('</pre>','')))

    #df['type'] = df['type']#.apply(lambda x: re.sub(r'[ ]+$', '', ''.join(x[0:-1]))) #.replace("'","|"))
    df = df.set_index('type')
    not_clusters = df.loc[misc]['percent'].sum()
    factor = (100 - not_clusters)/100
    df['percent normalized'] = df['percent'].apply(lambda x: round(x*factor,2))
    df.loc[misc,'percent normalized'] = ''

    
    outdfs.append(df)

df = outdfs[0]
if len(outdfs) > 1:
    for other_df in outdfs[1:]:
        df = pd.merge(left=df, right=other_df, left_on='type', right_on='type')

#df.to_csv('chto.csv')

pd.set_option('display.max_rows', df.shape[0]+1)
for index, row in df.iterrows():
    key = [x for x in list(depth_dict.keys()) if x in index][0]
    df.at[index, 'depth'] = int(depth_dict[key])
df['depth'] = df['depth'].astype(int)

position = ['']*(df['depth'].max()+1)

for index, row in df.iterrows():
    item = row['percent']
    position[row['depth']:] = ['']*(len(position[row['depth']:]))
    position[row['depth']] = index.replace('|','').replace('-','').replace('   ','')
    for ind, it in enumerate(position):
        df.loc[index, f'sum % at level {ind}'] = it

for ind, row in df[[x for x in df.columns if '% at' in x]].iterrows():
    chunk = list(filter(None,row.to_list()))
    type_name = re.sub('^[ ]+','', re.sub('[ ]+$','', chunk[-1])).replace("'",'')
    if len(chunk) > 1:
        parent_name = re.sub('^[ ]+','', re.sub('[ ]+$','', chunk[-2])).replace("'",'')
    else:
        parent_name = ''
    df.at[ind, 'type_name'] = type_name
    df.at[ind, 'parent_name'] = parent_name
    
for depth_level in reversed([x for x in df.columns if 'sum % at' in x]):
    for item in [x for x in df[depth_level].unique() if x != '']:
        sub_df = df[df[depth_level]==item]
        sub_sum = sub_df['percent'].sum()
        sub_id = sub_df.index[0]
        df.loc[sub_df.index, {depth_level}] = ''
        df.loc[sub_id, {depth_level}] = sub_sum
        
df['sum percent'] = df[[x for x in df.columns if 'at level' in x]].replace('',0).max(axis=1)
df = df.drop([x for x in df.columns if 'at level' in x], axis=1)
df = df.drop('depth', axis=1)

import plotly.graph_objects as go

fig =px.sunburst(
    df,
    names='type_name',
    parents='parent_name',
    values='sum percent',
    color='parent_name',
    branchvalues="total",
    hover_data=['percent normalized', 'percent']
)

#fig.update_layout(margin = dict(t=0, l=0, r=0, b=0))

fig.write_image("fig0.png", width=1000, height=1000)
fig.show()

import pandas as pd
import networkx as nx


Graph_df = nx.from_pandas_edgelist(df, 'type_name', 'parent_name', ['sum percent'], create_using=nx.MultiDiGraph()).reverse()




naming = ' '.join(get_cyphers()[INDEX][0:2])

# Create a Pandas Excel writer using XlsxWriter as the engine.
writer = pd.ExcelWriter(f'rep_{INDEX}.xlsx', engine='xlsxwriter')

temp_df = df.drop(['parent_name','type_name'],axis=1)
# Convert the dataframe to an XlsxWriter Excel object.
temp_df.to_excel(writer, sheet_name='Sheet1', startrow=1, header=False, startcol=1, index=False)

# Get the xlsxwriter objects from the dataframe writer object.
workbook  = writer.book
worksheet = writer.sheets['Sheet1']

header_format = workbook.add_format({
    'bold': True,
    'text_wrap': False,
    'valign': 'center',
    'align': 'left',
    'border': 1})

for col_num, value in enumerate(temp_df.columns.values):
    worksheet.write(0, col_num + 1, value, header_format)
    
for row_num, value in enumerate(temp_df.index.values):
    value = re.sub('\s+\|$','',value)
    worksheet.write(row_num + 1, 0, value, header_format)

worksheet.conditional_format('A14:A31', {'type': 'text',
                                     'criteria': 'containing',
                                     'value' : '-', 'format': workbook.add_format({'bg_color': 'green','align': 'left'})})
worksheet.conditional_format('A32:A53', {'type': 'text',
                                     'criteria': 'containing',
                                     'value' : '-', 'format': workbook.add_format({'bg_color': 'yellow','align': 'left'})})
worksheet.conditional_format('A57:G57', {'type': 'text',
                                     'criteria': 'containing',
                                     'value' : '-', 'format': workbook.add_format({'bg_color': 'pink','align': 'left'})})
worksheet.conditional_format('A58:G75', {'type': 'text',
                                     'criteria': 'containing',
                                     'value' : '-', 'format': workbook.add_format({'bg_color': 'blue','align': 'left'})})

worksheet.set_column(0, 0, 40)
worksheet.set_column(5, 6, 15)

worksheet.insert_image('H1','fig0.png')

##SECOND

required_names = ['rDNA',
                 'mobile_element',
                  'Class_I',
                  'Ty1_copia',
                  'Ty3_gypsy',
                  'Class_II',
                  'satellite',
                  'LINE',
                 'organelle',
                 'Unclassified repeat (No evidence)',
                 'contamination']
plotted_names= ['LINE', 'rDNA', 'organelle', 'satellite', 'Ty1_copia', 'Ty3_gypsy', 'Class_II', 'Unclassified repeat (No evidence)']

required_index = [[y for y in df.index if x in y][0] for x in required_names]

renamer = dict(zip(required_index, required_names))

scnd_df = df.loc[required_index, 'sum percent'].rename(renamer, axis=1)
scnd_df.to_excel(writer, sheet_name='Sheet2', startrow=0, header=True, startcol=0, index=True)
worksheet2 = writer.sheets['Sheet2']

fig = px.pie(scnd_df[plotted_names].reset_index(), values='sum percent', names='type', title=f'Типизация ридов {naming} по RepeatExplorer2')
fig.write_image("fig1.png")

worksheet2.insert_image('C4','fig1.png')

optional = []
for i in ['LINE', 'rDNA', 'organelle', 'contamination']:
    if scnd_df.at[i] != 0:
        if optional:
            break
        else:
            optional.append('Также, найдены')
            break

context_list = ['LINE', 'rDNA', 'organelle', 'contamination']
disc_list = ['элементы типа LINE ({:,.2f}%)', 'rDNA ({:,.2f}%)', 'последовательности органнельной ДНК ({:,.2f}%)', 'неклассифицированные повторяющиеся последовательности ({:,.2f}%)', 'продукты загрязнения адаптером ({:,.2f}%)']
context_dict = dict(zip(context_list, disc_list))

for num, i in enumerate(['LINE', 'rDNA', 'organelle', 'contamination']):
    if scnd_df.at[i] != 0:
        optional.append(context_dict[i].format(scnd_df.at[i]))
        
        
core=f'''В результате анализа репитома образца {naming} были выявлены следующие типы повторяющихся последовательностей ДНК: {'{:,.2f}'.format(scnd_df.at['mobile_element'])}% мобильных элементов, из которых {'{:,.2f}'.format(scnd_df.at['Class_I'])}% относятся к ретротранспозонам, а {'{:,.2f}'.format(scnd_df.at['Class_II'])}% к ДНК транспозонам. Надсемейство Ty1_copia составило {'{:,.2f}'.format(scnd_df.at['Ty1_copia'])}%, а Ty3_gypsy - {'{:,.2f}'.format(scnd_df.at['Ty3_gypsy'])}%. Доля сателлитов составила {'{:,.2f}'.format(scnd_df.at['satellite'])}%.  Неклассифицированные повторяющиеся последовательности составили {'{:,.2f}'.format(scnd_df.at['Unclassified repeat (No evidence)'])}%'''

txt = core +'\n'+', '.join(optional) + '.'

worksheet2.write('A13',txt)
worksheet2.set_column(0, 2, 20)

##THIRD
out_set=set()
parents = ['Class_I','Ty1_copia','Ty3_gypsy','Class_II']
for parent in parents:
    desc = nx.descendants(Graph_df, parent)
    children_with_value = [x for x in desc if df[df['type_name']==x]['percent'][0]>0]
    out_set.update(children_with_value)

required_names = list(out_set)

required_index = [[y for y in df.index if x in y][0] for x in required_names]

renamer = dict(zip(required_index, required_names))

thrd_df = df.loc[required_index, 'percent'].rename(renamer, axis=1)
thrd_df.to_excel(writer, sheet_name='Sheet3', startrow=0, header=True, startcol=0, index=True)
worksheet3 = writer.sheets['Sheet3']

##TY1
for num, parent in enumerate(parents):
    num = num+2
    desc = nx.descendants(Graph_df, parent)
    children_with_value = [x for x in desc if df[df['type_name']==x]['percent'][0]>0]
    plotted_names = children_with_value
    fig = px.pie(thrd_df[plotted_names].reset_index(), values='percent', names='type', title=f'{naming}, {parent}')
    if not 'Class' in parent:
        fig.write_image(f"fig{num}.png")
        worksheet3.insert_image(f'C{(num-3)*30+4}',f'fig{num}.png')


worksheet3.conditional_format('A1:B7', {'type': 'text',
                                     'criteria': 'containing',
                                     'value' : '-', 'format': workbook.add_format({'bg_color': 'green','align': 'left'})})
worksheet3.conditional_format('A8:B12', {'type': 'text',
                                     'criteria': 'containing',
                                     'value' : '-', 'format': workbook.add_format({'bg_color': 'yellow','align': 'left'})})
worksheet3.conditional_format('A13:B13', {'type': 'text',
                                     'criteria': 'containing',
                                     'value' : '-', 'format': workbook.add_format({'bg_color': 'pink','align': 'left'})})
worksheet3.conditional_format('A14:B16', {'type': 'text',
                                     'criteria': 'containing',
                                     'value' : '-', 'format': workbook.add_format({'bg_color': 'blue','align': 'left'})})

worksheet3.set_column(0, 2, 20)

txt = []

txt.append('Среди мобильных элементов можно выделить два класса, с выделением в первом классе подклассов Ty1_copia и Ty3_gypsy.')
for num,parent in enumerate(['Class_I','Ty1_copia','Ty3_gypsy','Class_II']):
    desc = nx.descendants(Graph_df, parent)
    children_with_value = [x for x in desc if df[df['type_name']==x]['percent'][0]>0]
    if parent =='Class_I':
        children_with_value = set(children_with_value) - set(nx.descendants(Graph_df, 'Ty1_copia')) - set(nx.descendants(Graph_df, 'Ty3_gypsy'))
        children_with_value.update(['Ty1_copia','Ty3_gypsy'])
        children_with_value = list(children_with_value)
    chunks = [f'{x} ({"{:,.2f}".format(df[df["type_name"]==x]["percent"][0])}%)' if (x not in ['Ty1_copia','Ty3_gypsy']) else f'{x} ({"{:,.2f}".format(df[df["type_name"]==x]["sum percent"][0])}%)' for x in children_with_value]
    if df[df["type_name"]==parent]["percent"][0] != 0:
        chunks.append(f'Неидентифицированный {parent} ({"{:,.2f}".format(df[df["type_name"]==parent]["percent"][0])}%)')
    _class = 'Подкласс' if parent in ['Ty1_copia','Ty3_gypsy'] else 'Класс'
    txt.append(f'{_class} {parent} представлен элементами: ' + ', '.join(chunks)+'.')


txt = ' '.join(txt)

worksheet3.write('A20',txt)

writer.save()
